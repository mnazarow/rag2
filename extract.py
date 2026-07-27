"""
Извлечение текста и метаданных из файлов базы знаний.

Порядок предпочтений для PDF:
  1. PyMuPDF (fitz)  — быстро, даёт постраничный текст и разметку;
  2. pdftotext (poppler-utils) — CLI-фолбэк, есть почти везде;
  3. если текстового слоя нет (скан) — документ помечается needs_ocr=1
     и уходит в очередь OCR (см. README, раздел «сканы»).

По результатам аудита базы: ~20% PDF — сканы без текстового слоя
(в основном сертификаты и декларации), их без OCR не прочитать.
"""
from __future__ import annotations

import csv
import hashlib
import io
import re
import subprocess
import unicodedata
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path

import config

# ------------------------------------------------------------- метаданные ---

DOC_TYPE_MAP = {
    # Ключи и с номером папки, и без — имена папок в базе не всегда единообразны.
    "КАТАЛОГ": "КАТАЛОГ",
    "ПРАЙС ЛИСТ": "ПРАЙС-ЛИСТ",
    "ПРАЙС-ЛИСТ": "ПРАЙС-ЛИСТ",
    "РУКОВОДСТВО": "РУКОВОДСТВО",
    "ПАСПОРТ": "ПАСПОРТ",
    "СЕРТИФИКАТ": "СЕРТИФИКАТ",
    "ДЕКЛАРАЦИИ": "СЕРТИФИКАТ",
    "ОПРОСНЫЙ ЛИСТ": "ОПРОСНЫЙ ЛИСТ",
    "РЕФЕРЕНС": "РЕФЕРЕНС",
    "ПРОГРАММА ПОДБОРА": "ПРОГРАММА ПОДБОРА",
    "ЧЕРТЕЖИ DWG REVIT 3D": "ЧЕРТЁЖ",
    "ВИДЕО ФОТОМАТЕРИАЛЫ": "МЕДИА",
    "1КАТАЛОГ": "КАТАЛОГ",
    "2ПРАЙС_ЛИСТ": "ПРАЙС-ЛИСТ",
    "3РУКОВОДСТВО": "РУКОВОДСТВО",
    "4ПАСПОРТ": "ПАСПОРТ",
    "5СЕРТИФИКАТ": "СЕРТИФИКАТ",
    "6ОПРОСНЫЙ ЛИСТ": "ОПРОСНЫЙ ЛИСТ",
    "7РЕФЕРЕНС": "РЕФЕРЕНС",
    "8РЕФЕРЕНС": "РЕФЕРЕНС",
    "7ПРОГРАММА ПОДБОРА": "ПРОГРАММА ПОДБОРА",
    "9ЧЕРТЕЖИ_DWG_REVIT_3D": "ЧЕРТЁЖ",
    "10ВИДЕО_ФОТОМАТЕРИАЛЫ": "МЕДИА",
    "ДРУГОЕ": "ДРУГОЕ",
    "АРХИВ": "АРХИВ",
    "Архив": "АРХИВ",
}

def nfc(text: str | None) -> str | None:
    """
    macOS хранит имена файлов в форме NFD: «Й» = «И» + комбинирующая бреве.
    Без нормализации папка «2ПРАЙС_ЛИСТ» на диске не совпадает с такой же
    строкой в коде, и поиск по имени файла молча ничего не находит.
    """
    return unicodedata.normalize("NFC", text) if text else text


# Папки, содержимое которых заведомо неактуально.
ARCHIVE_MARKERS = ("архив", "archive", "старые", "устаревш", "old", "не актуальн")

_DATE_PATTERNS = [
    # 01.02.2026 / 03-03-2026 / 18_08_2025
    (re.compile(r"(?<!\d)(\d{2})[._\-](\d{2})[._\-](20\d{2})(?!\d)"), ("d", "m", "y")),
    # 2026-04-15
    (re.compile(r"(?<!\d)(20\d{2})[._\-](\d{2})[._\-](\d{2})(?!\d)"), ("y", "m", "d")),
    # price_300326 / price_150126  (ддммгг)
    (re.compile(r"(?<!\d)(\d{2})(\d{2})(\d{2})(?!\d)"), ("d", "m", "yy")),
    # просто год: «Каталог SPL 2023»
    (re.compile(r"(?<!\d)(20[12]\d)(?!\d)"), ("y",)),
]


def parse_date_from_name(name: str) -> str | None:
    """Достаёт дату документа из имени файла. Возвращает ISO-дату или None."""
    for rx, order in _DATE_PATTERNS:
        m = rx.search(name)
        if not m:
            continue
        try:
            parts = dict(zip(order, m.groups()))
            if "yy" in parts:
                year = 2000 + int(parts["yy"])
            elif "y" in parts:
                year = int(parts["y"])
            else:
                continue
            month = int(parts.get("m", 1))
            day = int(parts.get("d", 1))
            if not (1 <= month <= 12 and 1 <= day <= 31 and 2000 <= year <= date.today().year + 2):
                continue
            return date(year, month, day).isoformat()
        except (ValueError, KeyError):
            continue
    return None


def version_key(rel_path: Path, brand: str | None, doc_type: str | None) -> str:
    """
    Ключ «одной и той же сущности документа»: по нему определяем,
    что новый прайс вытесняет старый.

    Из имени вырезаются даты, номера версий и хвосты «(1)», «(2)», «_v2».
    """
    stem = rel_path.stem
    s = stem
    for rx, _ in _DATE_PATTERNS:
        s = rx.sub(" ", s)
    s = re.sub(r"\bv\s?\d+(\.\d+)*\b", " ", s, flags=re.I)
    s = re.sub(r"\(\d+\)", " ", s)
    s = re.sub(r"[_\-\s]+", " ", s).strip().lower()
    return f"{(brand or '').lower()}|{(doc_type or '').lower()}|{s}"


@dataclass
class Meta:
    section: str | None = None
    brand: str | None = None
    doc_type: str | None = None
    effective_date: str | None = None
    version_key: str = ""
    is_archive: bool = False


def path_meta(abs_path: Path) -> Meta:
    rel = abs_path.relative_to(config.KB_ROOT)
    parts = tuple(nfc(p) for p in rel.parts)
    section = parts[0] if len(parts) > 1 else None
    brand = parts[1] if len(parts) > 2 else None
    # У SPL третий уровень — товарная линейка: «SPL_Радиаторы (Зайцев Руслан)».
    # Она информативнее бренда, поэтому дописываем её и убираем имя ответственного.
    if brand and len(parts) > 3 and re.match(r"^[A-ZА-Я]{2,}[_ ]", parts[2]):
        line = re.sub(r"\s*\([^)]*\)\s*$", "", parts[2]).replace("_", " ").strip()
        if line and line.lower() != brand.lower():
            brand = f"{brand} / {line}"
    doc_type = None
    for p in parts[2:-1]:                # только папки-категории, без раздела и файла
        key = p.strip()
        if key in DOC_TYPE_MAP:
            doc_type = DOC_TYPE_MAP[key]
        elif re.match(r"^\d{1,2}[А-ЯЁA-Z_][А-ЯЁA-Z_ ]{3,}$", key):   # 1КАТАЛОГ, 2ПРАЙС_ЛИСТ...
            plain = re.sub(r"^\d+", "", key).replace("_", " ").strip()
            doc_type = DOC_TYPE_MAP.get(key) or DOC_TYPE_MAP.get(plain) or plain
    if doc_type is None and abs_path.suffix.lower() in config.TABLE_EXTENSIONS:
        doc_type = "ТАБЛИЦА"
    is_archive = any(marker in p.lower() for p in parts[:-1] for marker in ARCHIVE_MARKERS)
    eff = parse_date_from_name(nfc(abs_path.name)) or parse_date_from_name(str(rel))
    return Meta(section, brand, doc_type, eff,
                version_key(rel, brand, doc_type), is_archive)


# ------------------------------------------------------------- извлечение ---

@dataclass
class Extracted:
    pages: list[str] = field(default_factory=list)   # текст постранично
    n_pages: int = 0
    needs_ocr: bool = False
    error: str | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(self.pages)

    @property
    def n_chars(self) -> int:
        return sum(len(p) for p in self.pages)


# Файлы крупнее этого порога хэшируются частично: начало, конец и размер.
# В базе 64 ГБ чертежей DWG и 15 ГБ видео — полное чтение при каждом
# обходе занимало бы минуты и не давало ничего: вероятность совпадения
# начала, конца и размера у двух разных чертежей пренебрежимо мала.
QUICK_HASH_THRESHOLD = 64 * 1024 * 1024
QUICK_HASH_WINDOW = 4 * 1024 * 1024


def file_hash(path: Path, block: int = 1 << 20) -> str:
    h = hashlib.sha256()
    try:
        size = path.stat().st_size
    except OSError:
        size = 0
    if size > QUICK_HASH_THRESHOLD:
        h.update(str(size).encode())
        with path.open("rb") as fh:
            h.update(fh.read(QUICK_HASH_WINDOW))
            fh.seek(-QUICK_HASH_WINDOW, 2)
            h.update(fh.read(QUICK_HASH_WINDOW))
        return "q" + h.hexdigest()
    with path.open("rb") as fh:
        while chunk := fh.read(block):
            h.update(chunk)
    return h.hexdigest()


def _pdf_pymupdf(path: Path) -> Extracted | None:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None
    try:
        pages: list[str] = []
        with fitz.open(path) as doc:
            for page in doc:
                pages.append(page.get_text("text") or "")
        chars = sum(len(p) for p in pages)
        return Extracted(pages, len(pages), needs_ocr=chars / max(len(pages), 1) < 100)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"pymupdf: {exc}")


def _pdf_poppler(path: Path) -> Extracted:
    try:
        info = subprocess.run(["pdfinfo", str(path)], capture_output=True,
                              text=True, timeout=60).stdout
        n_pages = 0
        for line in info.splitlines():
            if line.startswith("Pages:"):
                n_pages = int(line.split()[1])
                break
        out = subprocess.run(["pdftotext", "-q", "-layout", str(path), "-"],
                             capture_output=True, text=True, timeout=300).stdout
        pages = out.split("\f")
        if pages and not pages[-1].strip():
            pages.pop()
        n_pages = n_pages or len(pages)
        chars = sum(len(p) for p in pages)
        return Extracted(pages, n_pages, needs_ocr=chars / max(n_pages, 1) < 100)
    except FileNotFoundError:
        return Extracted(error="не найден pdftotext (поставьте poppler-utils) и PyMuPDF")
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"poppler: {exc}")


def extract_pdf(path: Path) -> Extracted:
    res = _pdf_pymupdf(path)
    if res is not None and not res.error and res.n_chars > 0:
        return res
    return _pdf_poppler(path)


def extract_docx(path: Path) -> Extracted:
    try:
        import docx  # python-docx
    except ImportError:
        return _docx_raw(path)
    try:
        d = docx.Document(str(path))
        lines: list[str] = []
        for para in d.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style = (para.style.name or "").lower()
            lines.append(f"\n## {t}\n" if "heading" in style or "заголов" in style else t)
        for table in d.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                lines.append("\n" + "\n".join(rows) + "\n")
        return Extracted(["\n".join(lines)], 1)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"docx: {exc}")


def _docx_raw(path: Path) -> Extracted:
    """Фолбэк без python-docx: вытаскиваем текст прямо из XML."""
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", xml)
        return Extracted([re.sub(r"\n{3,}", "\n\n", text)], 1)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"docx-raw: {exc}")


def extract_pptx(path: Path) -> Extracted:
    try:
        with zipfile.ZipFile(path) as z:
            slides = sorted(n for n in z.namelist()
                            if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
            pages = []
            for name in slides:
                xml = z.read(name).decode("utf-8", "ignore")
                texts = re.findall(r"<a:t>(.*?)</a:t>", xml, flags=re.S)
                pages.append("\n".join(t.strip() for t in texts if t.strip()))
        return Extracted(pages, len(pages))
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"pptx: {exc}")


def extract_xlsx(path: Path, max_rows: int = 5000) -> Extracted:
    """Таблица → построчный текст. Для прайсов параллельно работает prices.py."""
    try:
        import openpyxl
    except ImportError:
        return Extracted(error="нет openpyxl")
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        pages = []
        for ws in wb.worksheets:
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    rows.append(f"... (усечено, всего строк: {ws.max_row})")
                    break
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                pages.append(f"# Лист: {ws.title}\n" + "\n".join(rows))
        wb.close()
        return Extracted(pages, len(pages))
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"xlsx: {exc}")


def extract_xls(path: Path) -> Extracted:
    try:
        import xlrd
    except ImportError:
        return Extracted(error="нет xlrd (нужен для старых .xls)")
    try:
        book = xlrd.open_workbook(str(path))
        pages = []
        for sheet in book.sheets():
            rows = []
            for r in range(sheet.nrows):
                cells = [str(c).strip() for c in sheet.row_values(r) if str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                pages.append(f"# Лист: {sheet.name}\n" + "\n".join(rows))
        return Extracted(pages, len(pages))
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"xls: {exc}")


def extract_csv(path: Path) -> Extracted:
    for enc in ("utf-8-sig", "cp1251", "utf-8"):
        try:
            text = path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        return Extracted(error="csv: не определилась кодировка")
    try:
        dialect = csv.Sniffer().sniff(text[:4000], delimiters=";,\t")
        rows = [" | ".join(r) for r in csv.reader(io.StringIO(text), dialect)]
    except Exception:  # noqa: BLE001
        rows = text.splitlines()
    return Extracted(["\n".join(rows)], 1)


def extract_txt(path: Path) -> Extracted:
    for enc in ("utf-8", "cp1251", "utf-16", "latin-1"):
        try:
            return Extracted([path.read_text(encoding=enc)], 1)
        except (UnicodeDecodeError, LookupError):
            continue
    return Extracted(error="txt: не определилась кодировка")


def extract_html(path: Path) -> Extracted:
    """
    Сохранённые целиком страницы порталов производителей.
    trafilatura отделяет содержимое от навигации и рекламы; если её нет —
    работает грубый фолбэк. favor_recall=True: на страницах с таблицами
    характеристик точный режим склонен выбрасывать половину таблицы.
    """
    raw = extract_txt(path)
    if raw.error:
        return raw
    try:
        import trafilatura
        text = trafilatura.extract(raw.text, favor_recall=True,
                                   include_tables=True, include_links=False)
        if text and len(text) > 200:
            return Extracted([text], 1)
    except ImportError:
        pass
    text = re.sub(r"(?is)<(script|style|nav|footer|header|noscript).*?</\1>", " ", raw.text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;?", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [ln.strip() for ln in text.splitlines()]
    return Extracted(["\n".join(ln for ln in lines if len(ln) > 2)], 1)


def extract_mhtml(path: Path) -> Extracted:
    """MHTML — MIME-контейнер: достаём html-часть и разбираем как страницу."""
    try:
        import email
        msg = email.message_from_bytes(path.read_bytes())
        for part in msg.walk():
            if part.get_content_type() == "text/html":
                html = part.get_payload(decode=True).decode(
                    part.get_content_charset() or "utf-8", "ignore")
                tmp = Extracted([html], 1)
                try:
                    import trafilatura
                    text = trafilatura.extract(html, favor_recall=True, include_tables=True)
                    if text:
                        return Extracted([text], 1)
                except ImportError:
                    pass
                text = re.sub(r"(?is)<(script|style).*?</\1>", " ", tmp.text)
                return Extracted([re.sub(r"<[^>]+>", " ", text)], 1)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"mhtml: {exc}")
    return Extracted(error="mhtml: html-часть не найдена")


def extract_msg(path: Path) -> Extracted:
    """Outlook .msg — в базе их 78 (переписка о повышении цен и т.п.)."""
    try:
        import extract_msg  # type: ignore
    except ImportError:
        # Грубый фолбэк: вытащить читаемые строки из OLE-контейнера.
        data = path.read_bytes()
        text = data.decode("utf-16-le", "ignore")
        text = re.sub(r"[^\S\n]{3,}", " ", text)
        text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\t")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return Extracted([text[:200_000]], 1) if len(text) > 200 else Extracted(
            error="msg: нужен extract-msg")
    try:
        m = extract_msg.Message(str(path))
        head = f"От: {m.sender}\nКому: {m.to}\nДата: {m.date}\nТема: {m.subject}\n\n"
        return Extracted([head + (m.body or "")], 1)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"msg: {exc}")


def extract_rtf(path: Path) -> Extracted:
    """RTF: снимаем управляющие последовательности, декодируем \'xx."""
    raw = extract_txt(path)
    if raw.error:
        return raw
    text = raw.text
    text = re.sub(r"\\'([0-9a-fA-F]{2})",
                  lambda m: bytes([int(m.group(1), 16)]).decode("cp1251", "ignore"), text)
    text = re.sub(r"\\par[d]?", "\n", text)
    text = re.sub(r"\{\\\*?[^{}]*\}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"[ \t]{2,}", " ", text)
    body = "\n".join(ln.strip() for ln in text.splitlines() if ln.strip())
    return Extracted([body], 1) if body else Extracted(error="rtf: пусто")


def extract_opendocument(path: Path) -> Extracted:
    """ODT, ODS, ODP — тот же ZIP с XML внутри, что и у OpenOffice."""
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("content.xml").decode("utf-8", "ignore")
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"opendocument: {exc}")
    xml = re.sub(r"</text:(p|h)>", "\n", xml)
    xml = re.sub(r"</table:table-row>", "\n", xml)
    xml = re.sub(r"</table:table-cell>", " | ", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    text = re.sub(r"&amp;", "&", text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return Extracted(["\n".join(lines)], 1)


def extract_eml(path: Path) -> Extracted:
    """Письмо в формате EML — стандартный разбор без сторонних библиотек."""
    try:
        import email
        from email import policy
        msg = email.message_from_bytes(path.read_bytes(), policy=policy.default)
        head = (f"От: {msg.get('From','')}\nКому: {msg.get('To','')}\n"
                f"Дата: {msg.get('Date','')}\nТема: {msg.get('Subject','')}\n\n")
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body += part.get_content()
                elif part.get_content_type() == "text/html" and not body:
                    html = part.get_content()
                    body += re.sub(r"<[^>]+>", " ", html)
        else:
            body = msg.get_content()
        attachments = [a.get_filename() for a in msg.iter_attachments()
                       if a.get_filename()] if hasattr(msg, "iter_attachments") else []
        if attachments:
            body += "\n\nВложения: " + "; ".join(attachments)
        return Extracted([head + str(body)], 1)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"eml: {exc}")


def extract_epub(path: Path) -> Extracted:
    try:
        with zipfile.ZipFile(path) as z:
            names = [n for n in z.namelist()
                     if n.lower().endswith((".xhtml", ".html", ".htm"))]
            pages = []
            for name in sorted(names)[:400]:
                html = z.read(name).decode("utf-8", "ignore")
                html = re.sub(r"(?is)<(script|style).*?</\1>", " ", html)
                text = re.sub(r"<[^>]+>", " ", html)
                text = re.sub(r"\s{2,}", " ", text).strip()
                if len(text) > 50:
                    pages.append(text)
        return Extracted(pages, len(pages))
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"epub: {exc}")


def extract_fb2(path: Path) -> Extracted:
    raw = extract_txt(path)
    if raw.error:
        return raw
    text = re.sub(r"<binary[^>]*>.*?</binary>", " ", raw.text, flags=re.S)
    text = re.sub(r"</p>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    lines = [ln.strip() for ln in text.splitlines() if len(ln.strip()) > 2]
    return Extracted(["\n".join(lines)], 1)


def extract_djvu(path: Path) -> Extracted:
    """DjVu: нужен djvutxt из пакета djvulibre; иначе документ идёт в очередь OCR."""
    import shutil as _sh
    if not _sh.which("djvutxt"):
        return Extracted([], 0, needs_ocr=True, error=None)
    try:
        out = subprocess.run(["djvutxt", str(path)], capture_output=True,
                             text=True, timeout=300).stdout
        pages = out.split("\f")
        return Extracted(pages, len(pages), needs_ocr=len(out) < 100)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"djvu: {exc}")


def extract_ical(path: Path) -> Extracted:
    raw = extract_txt(path)
    if raw.error:
        return raw
    keep = ("SUMMARY", "DESCRIPTION", "LOCATION", "DTSTART", "DTEND", "ORGANIZER")
    lines = [ln for ln in raw.text.splitlines() if ln.split(":")[0].split(";")[0] in keep]
    return Extracted(["\n".join(lines)], 1)


def extract_vsdx(path: Path) -> Extracted:
    """Схемы Visio: текст надписей лежит в XML страниц."""
    try:
        with zipfile.ZipFile(path) as z:
            names = [n for n in z.namelist() if n.startswith("visio/pages/") and n.endswith(".xml")]
            pages = []
            for name in sorted(names):
                xml = z.read(name).decode("utf-8", "ignore")
                texts = re.findall(r"<Text[^>]*>(.*?)</Text>", xml, flags=re.S)
                clean = [re.sub(r"<[^>]+>", "", t).strip() for t in texts]
                body = "\n".join(t for t in clean if t)
                if body:
                    pages.append(body)
        return Extracted(pages, len(pages))
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"vsdx: {exc}")


def extract_raw_photo(path: Path) -> Extracted:
    """
    Сырой снимок Canon и подобные. Полное декодирование дорого и бессмысленно,
    если рядом лежит готовый JPG. Достаём только служебные данные съёмки;
    описание кадра делается по встроенному превью в media.py.
    """
    twin = None
    for ext in (".jpg", ".jpeg", ".JPG"):
        candidate = path.with_suffix(ext)
        if candidate.exists():
            twin = candidate
            break
    note = f"Есть готовый снимок: {twin.name}" if twin else "Готового снимка рядом нет"
    return Extracted([f"Сырой снимок (RAW). {note}"], 1)


def extract_svg(path: Path) -> Extracted:
    """SVG — это XML: текст внутри схемы достаётся бесплатно."""
    raw = extract_txt(path)
    if raw.error:
        return raw
    texts = re.findall(r"<(?:text|title|desc)[^>]*>(.*?)</(?:text|title|desc)>",
                       raw.text, flags=re.S | re.I)
    clean = [re.sub(r"<[^>]+>", " ", t).strip() for t in texts]
    body = "\n".join(t for t in clean if t)
    return Extracted([body], 1) if body else Extracted([], 0, needs_ocr=True)


def extract_step(path: Path) -> Extracted:
    """
    STEP и IGES — текстовые форматы ISO. Заголовок содержит имя изделия,
    автора и дату: это достаётся без всякого CAD-движка.
    """
    try:
        with path.open("r", encoding="utf-8", errors="ignore") as fh:
            head = fh.read(200_000)
    except OSError as exc:
        return Extracted(error=f"step: {exc}")
    lines: list[str] = []
    m = re.search(r"FILE_NAME\s*\((.*?)\);", head, flags=re.S | re.I)
    if m:
        lines.append("Заголовок STEP: " + re.sub(r"\s+", " ", m.group(1))[:600])
    for name in set(re.findall(r"PRODUCT\s*\(\s*'([^']{2,80})'", head, flags=re.I)):
        lines.append(f"Изделие: {name}")
    # IGES: секция Start (столбец 73 = 'S')
    if not lines and path.suffix.lower() in (".igs", ".iges"):
        for line in head.splitlines()[:40]:
            if len(line) > 72 and line[72] == "S":
                lines.append(line[:72].strip())
    body = "\n".join(dict.fromkeys(lines))
    return Extracted([body], 1) if body else Extracted([], 0, needs_ocr=True)


def extract_dxf(path: Path) -> Extracted:
    """
    DXF — текстовый формат AutoCAD. Достаём содержимое TEXT/MTEXT и атрибуты
    блоков: там лежит основная надпись чертежа (модель, масса, материал).
    DWG бинарный — его нужно сначала конвертировать (см. cadtools.py).
    """
    raw = extract_txt(path)
    if raw.error:
        return raw
    out, take = [], False
    for line in raw.text.splitlines():
        code = line.strip()
        if code in ("1", "3", "2"):        # текстовые группы DXF
            take = True
            continue
        if take:
            val = line.strip()
            take = False
            if 2 < len(val) < 200 and not val.startswith(("AcDb", "{\\f", "*")):
                out.append(val)
    uniq = list(dict.fromkeys(out))
    body = "\n".join(uniq)
    return Extracted([body], 1) if body else Extracted([], 0, needs_ocr=True)


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": _docx_raw,
    ".pptx": extract_pptx,
    ".xlsx": extract_xlsx,
    ".xlsm": extract_xlsx,
    ".xls": extract_xls,
    ".csv": extract_csv,
    ".txt": extract_txt,
    ".md": extract_txt,
    ".html": extract_html,
    ".htm": extract_html,
    ".mhtml": extract_mhtml,
    ".mht": extract_mhtml,
    ".msg": extract_msg,
    ".svg": extract_svg,
    ".stp": extract_step,
    ".step": extract_step,
    ".igs": extract_step,
    ".iges": extract_step,
    ".dxf": extract_dxf,
    ".json": extract_txt,
    ".xml": extract_txt,
    ".rtf": extract_rtf,
    ".odt": extract_opendocument,
    ".ods": extract_opendocument,
    ".odp": extract_opendocument,
    ".ott": extract_opendocument,
    ".eml": extract_eml,
    ".epub": extract_epub,
    ".fb2": extract_fb2,
    ".djvu": extract_djvu,
    ".djv": extract_djvu,
    ".ics": extract_ical,
    ".vsdx": extract_vsdx,
    ".vsd": extract_vsdx,
    ".cr2": extract_raw_photo,
    ".nef": extract_raw_photo,
    ".arw": extract_raw_photo,
    ".dng": extract_raw_photo,
    ".tsv": extract_csv,
    ".yaml": extract_txt,
    ".yml": extract_txt,
    ".ini": extract_txt,
    ".cfg": extract_txt,
    ".conf": extract_txt,
    ".log": extract_txt,
    ".sql": extract_txt,
    ".url": extract_txt,
    ".webloc": extract_txt,
}


def extract(path: Path) -> Extracted:
    fn = EXTRACTORS.get(path.suffix.lower())
    if fn is None:
        return Extracted(error=f"нет обработчика для {path.suffix}")
    try:
        return fn(path)
    except Exception as exc:  # noqa: BLE001
        return Extracted(error=f"{type(exc).__name__}: {exc}")


# ------------------------------------------------- карточки объектов --------

def _neighbour_context(path: Path, limit: int = 12) -> list[str]:
    """Имена соседних файлов: они описывают, о чём эта папка."""
    try:
        names = [nfc(p.name) for p in path.parent.iterdir()
                 if p.is_file() and p.name != path.name][:limit]
    except OSError:
        return []
    return [n for n in names if n and not n.startswith(config.IGNORE_PREFIXES)]


ASSET_KIND_RU = {
    "image": "изображение",
    "video": "видеозапись",
    "audio": "аудиозапись",
    "drawing": "чертёж / 3D-модель",
    "archive": "архив",
    "other": "файл",
}


def asset_kind(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in config.IMAGE_EXTENSIONS:
        return "image"
    if ext in config.VIDEO_EXTENSIONS:
        return "video"
    if ext in config.AUDIO_EXTENSIONS:
        return "audio"
    if ext in config.DRAWING_EXTENSIONS:
        return "drawing"
    if ext in config.ARCHIVE_EXTENSIONS:
        return "archive"
    return "other"


def asset_card(path: Path, meta: "Meta", description: str = "",
               transcript_hint: str = "") -> Extracted:
    """
    Текстовая карточка нетекстового файла.

    Смысл: DWG-чертёж, фотография или видео сами по себе для поиска невидимы.
    Карточка описывает объект словами — что это, какого бренда, из какой
    товарной линейки, как называется файл, что лежит рядом — и тем самым
    «привязывает» его к вышестоящим папкам. Если подключены распознавание
    речи и описание изображений, их результат дописывается сюда же.
    """
    kind = asset_kind(path)
    name = nfc(path.stem) or ""
    readable = re.sub(r"[_\-]+", " ", name).strip()
    lines = [f"{ASSET_KIND_RU[kind].capitalize()}: {readable}",
             f"Имя файла: {nfc(path.name)}"]
    if meta.brand:
        lines.append(f"Бренд / линейка: {meta.brand}")
    if meta.section:
        lines.append(f"Раздел базы: {meta.section}")
    if meta.doc_type:
        lines.append(f"Категория: {meta.doc_type}")
    try:
        rel_dir = path.parent.relative_to(config.KB_ROOT)
        lines.append("Расположение: " + " / ".join(nfc(x) for x in rel_dir.parts))
    except ValueError:
        pass
    if meta.effective_date:
        lines.append(f"Дата: {meta.effective_date}")
    try:
        size_mb = path.stat().st_size / 1048576
        lines.append(f"Размер: {size_mb:.1f} МБ")
    except OSError:
        pass
    if kind == "video":
        info = video_info(path)
        if info:
            lines.append(info)
    neighbours = _neighbour_context(path)
    if neighbours:
        lines.append("Рядом в папке: " + "; ".join(neighbours))
    if description:
        lines.append("Описание: " + description)
    if transcript_hint:
        lines.append("Расшифровка: " + transcript_hint)
    return Extracted(["\n".join(lines)], 1)


def video_info(path: Path) -> str:
    """Длительность и разрешение через ffprobe, если он установлен."""
    try:
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries",
             "format=duration:stream=width,height,codec_type",
             "-of", "default=nw=1", str(path)],
            capture_output=True, text=True, timeout=30).stdout
    except (FileNotFoundError, subprocess.SubprocessError):
        return ""
    dur = re.search(r"duration=([\d.]+)", out)
    wh = re.search(r"width=(\d+)\s+height=(\d+)", out)
    bits = []
    if dur:
        seconds = float(dur.group(1))
        bits.append(f"длительность {int(seconds // 60)} мин {int(seconds % 60)} с")
    if wh:
        bits.append(f"{wh.group(1)}×{wh.group(2)}")
    if "codec_type=audio" in out:
        bits.append("со звуковой дорожкой")
    return ("Параметры: " + ", ".join(bits)) if bits else ""


# ----------------------------------------------------------- архивы ---------

def unpack_archive(path: Path, dest: Path, depth: int = 0) -> tuple[int, str | None]:
    """
    Распаковывает архив в dest. Возвращает (сколько файлов, ошибка).

    Внутри архивов базы лежат сертификаты, старые паспорта и Revit-модели —
    без распаковки они выпадают из индекса целиком.
    """
    if depth > config.ARCHIVE_MAX_DEPTH:
        return 0, "превышена глубина вложенности архивов"
    dest.mkdir(parents=True, exist_ok=True)
    ext = path.suffix.lower()
    try:
        if ext == ".zip":
            with zipfile.ZipFile(path) as z:
                total = sum(i.file_size for i in z.infolist())
                if path.stat().st_size and total / path.stat().st_size > config.ARCHIVE_MAX_RATIO:
                    return 0, "подозрительная степень сжатия (защита от zip-бомбы)"
                z.extractall(dest)
        else:
            tool = None
            for candidate in (["bsdtar", "-xf"], ["7z", "x", "-y", "-o" + str(dest)],
                              ["unar", "-o", str(dest)], ["unrar", "x", "-y"]):
                import shutil as _sh
                if _sh.which(candidate[0]):
                    tool = candidate
                    break
            if tool is None:
                return 0, "нет распаковщика (поставьте bsdtar / p7zip / unar / unrar)"
            cmd = list(tool)
            if tool[0] == "bsdtar":
                cmd += [str(path), "-C", str(dest)]
            elif tool[0] == "unrar":
                cmd += [str(path), str(dest) + "/"]
            else:
                cmd += [str(path)]
            res = subprocess.run(cmd, capture_output=True, text=True, timeout=900)
            if res.returncode != 0:
                return 0, f"распаковщик вернул ошибку: {res.stderr[:150]}"
    except Exception as exc:  # noqa: BLE001
        return 0, f"архив: {exc}"
    files = [f for f in dest.rglob("*") if f.is_file()]
    # Вложенные архивы — рекурсивно.
    for inner in list(files):
        if inner.suffix.lower() in config.ARCHIVE_EXTENSIONS:
            unpack_archive(inner, inner.with_suffix(inner.suffix + ".unpacked"), depth + 1)
    return len(files), None


# ------------------------------------------------------------ фильтрация ----

def classify(path: Path) -> tuple[str, str]:
    """
    Что делать с файлом: text | table | asset | archive | skip.
    Второй элемент — причина пропуска.
    """
    name = path.name
    if name in config.IGNORE_NAMES or name.startswith(config.IGNORE_PREFIXES):
        return "skip", "служебный файл"
    ext = path.suffix.lower()
    if any(part.endswith(tuple(config.IGNORE_DIR_PARTS)) for part in path.parts):
        return "skip", "ресурсы сохранённой веб-страницы"
    if ext in config.IGNORE_EXTENSIONS:
        return "skip", "нерелевантное расширение"
    try:
        if path.stat().st_size > config.MAX_FILE_MB * 1024 * 1024 and \
                ext not in config.VIDEO_EXTENSIONS and ext not in config.ARCHIVE_EXTENSIONS:
            return "skip", f"больше {config.MAX_FILE_MB} МБ"
    except OSError:
        return "skip", "недоступен"
    if ext in config.ARCHIVE_EXTENSIONS:
        return ("archive", "") if config.EXTRACT_ARCHIVES else ("asset", "")
    if ext in config.TABLE_EXTENSIONS:
        try:
            if path.stat().st_size > config.TABLE_MAX_MB * 1024 * 1024:
                return "asset", ""      # программа подбора, а не прайс
        except OSError:
            pass
        return "table", ""
    if ext in EXTRACTORS:
        return "text", ""
    if ext in config.IMAGE_EXTENSIONS or ext in config.VIDEO_EXTENSIONS \
            or ext in config.AUDIO_EXTENSIONS or ext in config.DRAWING_EXTENSIONS:
        return ("asset", "") if config.ASSET_CARDS else ("skip", "нетекстовый формат")
    return "skip", "неизвестный формат"


def should_index(path: Path) -> tuple[bool, str]:
    """Совместимость со старым интерфейсом: обёртка над classify()."""
    kind, reason = classify(path)
    return kind != "skip", reason
